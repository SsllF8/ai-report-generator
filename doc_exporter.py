import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def export_to_word(report_data: dict, output_dir: str = "reports") -> str:
    """
    将结构化报告数据导出为 Word 文档
    
    Args:
        report_data: 报告数据字典
        output_dir: 输出目录
    
    Returns:
        str: 导出文件的路径
    """
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    
    # ============ 页面设置 ============
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # ============ 标题 ============
    template_icon = report_data.get("_template_icon", "")
    template_name = report_data.get("_template_name", "报告")
    title_text = report_data.get("title", f"{template_name}")
    
    # 主标题
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{template_icon} {title_text}")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%Y 年 %m 月 %d 日"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    date_run.font.name = "微软雅黑"
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    
    doc.add_paragraph()  # 空行
    
    # ============ 摘要 ============
    if report_data.get("summary") and not report_data.get("error"):
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"摘要：{report_data['summary']}")
        summary_run.font.size = Pt(11)
        summary_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        summary_run.font.name = "微软雅黑"
        summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        summary_run.italic = True
        doc.add_paragraph()
    
    # ============ 错误提示 ============
    if report_data.get("error"):
        error_para = doc.add_paragraph()
        error_run = error_para.add_run(f"生成出错：{report_data['error']}")
        error_run.font.size = Pt(12)
        error_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        error_run.font.name = "微软雅黑"
        error_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    
    # ============ 各段落内容 ============
    # 定义各模板对应的段落映射
    section_mapping = {
        "tech_weekly": [
            ("completed", "本周完成"),
            ("in_progress", "进行中"),
            ("issues", "问题与风险"),
            ("next_week", "下周计划"),
        ],
        "project_weekly": [
            ("progress", "项目进展"),
            ("milestones", "里程碑"),
            ("risks", "风险预警"),
            ("resource_needs", "资源需求"),
        ],
        "work_summary": [
            ("achievements", "工作成果"),
            ("growth", "能力成长"),
            ("improvements", "不足与改进"),
            ("future_plans", "未来规划"),
        ],
        "meeting_notes": [
            ("discussions", "议题讨论"),
            ("decisions", "决议事项"),
            ("action_items", "行动项"),
        ],
    }
    
    template_key = report_data.get("_template_key", "tech_weekly")
    sections = section_mapping.get(template_key, section_mapping["tech_weekly"])
    
    for field_name, section_title in sections:
        items = report_data.get(field_name, [])
        if not items:
            continue
        
        # 段落标题
        heading = doc.add_heading(level=1)
        run = heading.add_run(section_title)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        
        # 添加分隔线效果（用底部边框的空段落）
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(6)
        pPr = sep._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '6',
            qn('w:space'): '1',
            qn('w:color'): '3498DB'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        # 内容列表
        for i, item in enumerate(items):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_after = Pt(4)
            
            # 序号圆点
            bullet_run = para.add_run(f"  ●  ")
            bullet_run.font.size = Pt(10)
            bullet_run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
            
            # 内容
            text_run = para.add_run(str(item))
            text_run.font.size = Pt(11)
            text_run.font.name = "微软雅黑"
            text_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            text_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        doc.add_paragraph()  # 段落间空行
    
    # ============ 页脚 ============
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"由 AI 报告生成器自动生成  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
    footer_run.font.name = "微软雅黑"
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    
    # ============ 保存 ============
    # 清理文件名中的非法字符
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title_text)
    filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filepath
